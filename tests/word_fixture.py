"""Builds the Word fixture from code rather than checking in a binary.

A `.docx` in the repository is a blob nobody can review. This module says in
Python exactly what the fixture contains, which is the part that matters: the
heading styles a publisher template renames, a Zotero field citation, a
coauthor's tracked insertion, a commented range, a footnote, a table, and an
equation.

Two of those need raw XML because python-docx has no API for them. Fields and
tracked changes are written as elements, and the footnotes part is added to the
package by hand, because a footnote is where a surprising number of citations
live and a fixture without one would not exercise the path.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
M = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'

ZOTERO_ITEM = {
    "citationItems": [
        {
            "itemData": {
                "id": "nakamura2023",
                "title": "Budget-Matched Comparison of Sparse and Dense Retrieval",
                "author": [{"family": "Nakamura", "given": "H"}],
                "issued": {"date-parts": [["2023"]]},
                "DOI": "10.1145/3539618.3591998",
            }
        }
    ]
}

CLAIM = (
    "We present a comparison that holds the retrieval budget fixed rather than "
    "the index size, and we report recall at ten for every configuration we "
    "measured across the three corpora in this study."
)


def _xml(fragment: str) -> Any:
    from docx.oxml import parse_xml

    return parse_xml(fragment)


def _heading_style(document: Any, name: str, outline: int) -> str:
    """A style whose name says nothing about headings, only its outline level.

    This is the case a publisher template creates. The name is renamed to suit
    the publisher and the outline level is left alone, because Word builds its
    table of contents from the outline level.
    """
    from docx.enum.style import WD_STYLE_TYPE

    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    properties = style.element.get_or_add_pPr()
    properties.append(_xml(f'<w:outlineLvl {W} w:val="{outline}"/>'))
    return name


def _field_citation(paragraph: Any, rendered: str, payload: dict[str, Any]) -> None:
    instruction = f"ADDIN ZOTERO_ITEM CSL_CITATION {json.dumps(payload)}"
    paragraph._p.append(_xml(f'<w:r {W}><w:fldChar w:fldCharType="begin"/></w:r>'))
    paragraph._p.append(
        _xml(f'<w:r {W}><w:instrText xml:space="preserve">{instruction}</w:instrText></w:r>')
    )
    paragraph._p.append(_xml(f'<w:r {W}><w:fldChar w:fldCharType="separate"/></w:r>'))
    paragraph._p.append(_xml(f'<w:r {W}><w:t xml:space="preserve">{rendered}</w:t></w:r>'))
    paragraph._p.append(_xml(f'<w:r {W}><w:fldChar w:fldCharType="end"/></w:r>'))


def _tracked_insertion(paragraph: Any, text: str, author: str) -> None:
    paragraph._p.append(
        _xml(
            f'<w:ins {W} w:id="900" w:author="{author}" w:date="2026-02-01T00:00:00Z">'
            f'<w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:ins>'
        )
    )


def _tracked_deletion(paragraph: Any, text: str, author: str) -> None:
    """Text a coauthor has already struck out. It is not in the document."""
    paragraph._p.append(
        _xml(
            f'<w:del {W} w:id="901" w:author="{author}" w:date="2026-02-01T00:00:00Z">'
            f'<w:r><w:delText xml:space="preserve">{text}</w:delText></w:r></w:del>'
        )
    )


def _commented(paragraph: Any, text: str, comment_id: str = "1") -> None:
    paragraph._p.append(_xml(f'<w:commentRangeStart {W} w:id="{comment_id}"/>'))
    paragraph._p.append(_xml(f'<w:r {W}><w:t xml:space="preserve">{text}</w:t></w:r>'))
    paragraph._p.append(_xml(f'<w:commentRangeEnd {W} w:id="{comment_id}"/>'))


def _equation(paragraph: Any, text: str) -> None:
    paragraph._p.append(_xml(f"<m:oMath {M}><m:r><m:t>{text}</m:t></m:r></m:oMath>"))


def _add_footnotes(document: Any, notes: list[str]) -> None:
    from docx.opc.constants import CONTENT_TYPE, RELATIONSHIP_TYPE
    from docx.opc.packuri import PackURI
    from docx.opc.part import Part

    body = "".join(
        f'<w:footnote w:id="{index + 2}"><w:p><w:r>'
        f'<w:t xml:space="preserve">{text}</w:t></w:r></w:p></w:footnote>'
        for index, text in enumerate(notes)
    )
    blob = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:footnotes {W}>"
        '<w:footnote w:id="0" w:type="separator"><w:p><w:r><w:t> </w:t></w:r></w:p></w:footnote>'
        f"{body}</w:footnotes>"
    ).encode()

    part = Part(
        PackURI("/word/footnotes.xml"),
        CONTENT_TYPE.WML_FOOTNOTES,
        blob,
        document.part.package,
    )
    document.part.relate_to(part, RELATIONSHIP_TYPE.FOOTNOTES)


def build(target: Path) -> Path:
    """Write the fixture and return its path."""
    import docx

    document = docx.Document()
    document.core_properties.title = "Sparse Retrieval at Equal Cost"
    document.core_properties.author = "A. Researcher"

    chapter = _heading_style(document, "Chapter Title", outline=0)
    subhead = _heading_style(document, "Sub Head", outline=1)

    document.add_paragraph("Introduction", style=chapter)
    document.add_paragraph(
        "It is important to note that dense encoders are reported to beat sparse "
        "baselines on this benchmark, and the comparisons behind that claim hold "
        "the index size constant rather than the budget."
    )
    document.add_paragraph(CLAIM)

    cited = document.add_paragraph("The gap closes once the cost is held fixed ")
    _field_citation(cited, "(Nakamura, 2023)", ZOTERO_ITEM)
    cited.add_run(", and the effect holds at every budget we tried.")

    document.add_paragraph("Method", style=chapter)
    document.add_paragraph("Indexing", style=subhead)
    method = document.add_paragraph(
        "The corpus is indexed with BM25 and queries are expanded with the top "
        "three terms from a first pass over the corpus. "
    )
    _equation(method, "s(q,d) = sum idf(t)")

    coauthored = document.add_paragraph("Needless to say, the expansion is applied first. ")
    _tracked_insertion(coauthored, "This sentence was added by a coauthor. ", "B. Coauthor")
    _tracked_deletion(coauthored, "This sentence was struck out by a coauthor. ", "B. Coauthor")
    _commented(coauthored, "This clause is already under discussion.")

    document.add_paragraph("Results", style="Heading 1")
    document.add_paragraph(
        "Recall at ten rises from 0.62 to 0.71 with expansion enabled, and the "
        "cost is one third of the dense baseline over the five thousand queries "
        "we sampled for this comparison."
    )
    document.add_paragraph("We significantly outperform the dense encoder at the budget we fixed.")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "System"
    table.cell(0, 1).text = "Recall at ten"
    table.cell(1, 0).text = "BM25 with expansion"
    table.cell(1, 1).text = "0.71"

    document.add_paragraph("References", style=chapter)
    document.add_paragraph(
        "[1] H. Nakamura, Budget-Matched Comparison of Sparse and Dense Retrieval, 2023. "
        "doi:10.1145/3539618.3591998"
    )
    document.add_paragraph(
        "[2] J. Smith, A Study That Does Not Exist, Journal of Invented Results, 2021."
    )

    _add_footnotes(document, ["Recall at ten is reported throughout, see [1] for the protocol."])

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target))
    return target
