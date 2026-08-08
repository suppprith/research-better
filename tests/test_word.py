"""Word ingest, and writeback as tracked changes.

The fixture is built from code rather than checked in, so what it contains is
readable. See word_fixture.py.

Two things carry the most weight. Heading detection has to survive a publisher
template renaming every style, because that is the file a Springer or Elsevier
author actually has. And the tracked changes have to accept and reject cleanly,
because a change Word cannot reject is a change the author cannot decline.
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Any

import pytest

from research_better.edit import writeback
from research_better.edit.word import AUTHOR, WordWritebackError
from research_better.ingest import load, supported_suffixes
from research_better.ingest.word import W, extract, heading_level
from research_better.model import Document, FloatKind
from test_writeback import ledger_for
from word_fixture import CLAIM, build


@pytest.fixture(scope="module")
def source(tmp_path_factory: pytest.TempPathFactory) -> Path:
    return build(tmp_path_factory.mktemp("word-source") / "paper.docx")


@pytest.fixture
def draft(source: Path, tmp_path: Path) -> Path:
    target = tmp_path / "paper.docx"
    target.write_bytes(source.read_bytes())
    return target


@pytest.fixture
def document(draft: Path) -> Document:
    return load(draft)


# Dispatch ------------------------------------------------------------------


def test_docx_is_a_supported_format() -> None:
    assert ".docx" in supported_suffixes()


def test_a_word_draft_ingests_through_the_normal_entry_point(document: Document) -> None:
    assert document.format == "word"
    assert document.sentences


def test_the_extracted_text_is_the_source_text(draft: Path, document: Document) -> None:
    # Word is the first format with no linear source. The extracted prose is
    # what every offset in the model refers to, so the two have to be the same
    # string or every span in every artifact points at nothing.
    assert extract(draft).text == document.source_text


# Sections ------------------------------------------------------------------


def test_a_renamed_template_style_is_still_a_heading(document: Document) -> None:
    # A publisher template renames the heading styles and leaves the outline
    # level alone, because Word builds its table of contents from the outline
    # level. That is the signal a template cannot rename away.
    titles = [(section.title, section.level) for section in document.sections]
    assert ("Introduction", 1) in titles
    assert ("Indexing", 2) in titles


def test_a_built_in_heading_style_still_works(document: Document) -> None:
    assert ("Results", 1) in [(section.title, section.level) for section in document.sections]


def test_the_section_tree_nests(document: Document) -> None:
    indexing = next(section for section in document.sections if section.title == "Indexing")
    assert indexing.path == ("Method", "Indexing")


def test_a_style_with_no_outline_level_or_heading_name_is_body_text() -> None:
    assert heading_level("BodyText", {}) is None
    assert heading_level(None, {}) is None


# Citations -----------------------------------------------------------------


def test_a_field_citation_resolves_to_metadata_not_the_rendered_string(
    document: Document,
) -> None:
    field = next(c for c in document.citations if c.key == "nakamura2023")
    assert field.resolved is not None
    # The reference manager knows the DOI. Parsing "(Nakamura, 2023)" never
    # would, and that string is all a text-only reader would see.
    assert field.resolved.doi == "10.1145/3539618.3591998"
    assert field.resolved.title.startswith("Budget-Matched Comparison")
    assert field.resolved.year == 2023
    assert field.resolved.source == "reference manager"


def test_a_field_citation_is_attached_to_the_sentence_it_sits_in(document: Document) -> None:
    field = next(c for c in document.citations if c.key == "nakamura2023")
    assert field.sentence_id
    assert "The gap closes" in document.sentence(field.sentence_id).text


def test_the_rendered_citation_is_not_editable(document: Document) -> None:
    field = next(c for c in document.citations if c.key == "nakamura2023")
    # Word refreshes a field from its instruction. An edit inside the rendered
    # text would be silently overwritten the next time it did.
    assert not document.is_patchable(field.span)


def test_the_reference_list_becomes_bibliography_entries(document: Document) -> None:
    entries = [c.key for c in document.citations if c.in_bibliography]
    assert entries == ["1", "2"]


def test_a_field_instruction_that_cannot_be_read_yields_no_citation() -> None:
    from research_better.ingest.word import _citation_from_instruction

    # Guessing at a citation from a half-parsed instruction would put an
    # invented record into the bibliography.
    assert _citation_from_instruction("ADDIN ZOTERO_ITEM CSL_CITATION {not json") == []
    assert _citation_from_instruction("PAGEREF _Toc123") == []


# What is not prose ---------------------------------------------------------


def test_a_table_is_located_and_never_segmented(document: Document) -> None:
    table = next(item for item in document.floats if item.kind is FloatKind.TABLE)
    body = document.text_of(table.span)
    assert "BM25 with expansion" in body
    assert not any(sentence.span.overlaps(table.span) for sentence in document.sentences)


def test_an_equation_is_opaque_to_segmentation_and_to_patching(document: Document) -> None:
    equation = next(item for item in document.floats if item.kind is FloatKind.EQUATION)
    assert not document.is_patchable(equation.span)


def test_a_footnote_is_captured_but_is_not_body_prose(document: Document) -> None:
    note = next(item for item in document.floats if (item.label or "").startswith("footnote"))
    # A footnote is where a surprising number of citations live, so it is
    # extracted. It is not a sentence in the argument, so it is not segmented.
    assert "reported throughout" in document.text_of(note.span)
    assert not any(sentence.span.overlaps(note.span) for sentence in document.sentences)


# A coauthor's work ---------------------------------------------------------


def test_text_a_coauthor_deleted_is_not_extracted(document: Document) -> None:
    assert "struck out by a coauthor" not in document.source_text


def test_text_a_coauthor_inserted_is_extracted_and_left_alone(document: Document) -> None:
    start = document.source_text.index("This sentence was added by a coauthor")
    assert not document.is_patchable(_span(document, start, start + 10))


def test_a_commented_range_is_left_alone(document: Document) -> None:
    start = document.source_text.index("This clause is already under discussion")
    # A comment is a conversation in progress. The tool does not get to edit
    # inside one.
    assert not document.is_patchable(_span(document, start, start + 10))


def test_the_claim_is_still_found(document: Document) -> None:
    from research_better import novelty

    assert novelty.analyse(document).claim == " ".join(CLAIM.split())


# Writeback -----------------------------------------------------------------


def test_writeback_produces_tracked_changes(draft: Path, document: Document) -> None:
    rows = list(ledger_for(document).edits)
    assert rows, "an empty ledger would make this test prove nothing"

    writeback.apply(document, rows, force=True)

    body = _body(draft)
    ours = [
        element
        for element in list(body.iter(f"{W}ins")) + list(body.iter(f"{W}del"))
        if element.get(f"{W}author") == AUTHOR
    ]
    assert ours, "the edits have to arrive in the review pane, not as silent replacements"
    for element in ours:
        assert element.get(f"{W}date")


def test_a_coauthors_tracked_change_survives_writeback(draft: Path, document: Document) -> None:
    writeback.apply(document, list(ledger_for(document).edits), force=True)
    body = _body(draft)
    authors = {element.get(f"{W}author") for element in body.iter(f"{W}ins")}
    assert "B. Coauthor" in authors


def test_accepting_every_change_gives_the_edited_paper(draft: Path, document: Document) -> None:
    rows = list(ledger_for(document).edits)
    expected = _apply_offline(document.source_text, rows)

    writeback.apply(document, rows, force=True)
    _resolve(draft, accept=True)
    assert extract(draft).text == expected


def test_rejecting_every_change_gives_the_paper_back(draft: Path, document: Document) -> None:
    before = document.source_text
    writeback.apply(document, list(ledger_for(document).edits), force=True)

    _resolve(draft, accept=False)
    # A change Word cannot reject is a change the author cannot decline.
    assert extract(draft).text == before


def test_writeback_refuses_a_file_that_has_moved_on(draft: Path, document: Document) -> None:
    import docx

    package = docx.Document(str(draft))
    package.add_paragraph("Somebody typed a sentence while the analysis was running.")
    package.save(str(draft))

    with pytest.raises(WordWritebackError, match="no longer extracts"):
        writeback.apply(document, list(ledger_for(document).edits), force=True)


def test_writeback_takes_a_backup(draft: Path, document: Document) -> None:
    original = draft.read_bytes()
    written = writeback.apply(document, list(ledger_for(document).edits), force=True)
    assert written.backups[0].read_bytes() == original


def test_the_edit_command_writes_a_word_draft(draft: Path) -> None:
    from research_better.artifacts import ArtifactStore
    from research_better.cli import EXIT_FINDINGS, main

    for command in ("novelty", "ground", "fluff", "voice"):
        arguments = [command, str(draft), "--quiet"]
        if command == "novelty":
            arguments.insert(2, "--confirm-claim")
        main(arguments)

    assert main(["edit", str(draft), "--apply", "--quiet"]) == EXIT_FINDINGS
    artifact = ArtifactStore(draft).read("edits")
    assert artifact is not None
    assert artifact.payload["written"]["files"]


# Helpers -------------------------------------------------------------------


def _span(document: Document, start: int, end: int) -> Any:
    from research_better.model import Span

    return Span(start, end)


def _body(path: Path) -> Any:
    import docx

    return docx.Document(str(path)).element.body


def _apply_offline(text: str, rows: list[Any]) -> str:
    from research_better.edit.ledger import apply_to

    return apply_to(text, rows)


def _resolve(path: Path, accept: bool) -> None:
    """Do what Word's accept-all or reject-all does, for our changes only.

    Accepting drops the deletions and keeps the insertions. Rejecting does the
    reverse and turns the deleted text back into ordinary text. A coauthor's
    changes are left as they are, because a test that resolved those too would
    not be testing what an author sees when they click accept on ours.
    """
    import docx

    package = docx.Document(str(path))
    body = package.element.body

    for element in list(body.iter(f"{W}ins")) + list(body.iter(f"{W}del")):
        if element.get(f"{W}author") != AUTHOR:
            continue
        insertion = element.tag == f"{W}ins"
        keep = insertion if accept else not insertion
        parent = element.getparent()
        position = list(parent).index(element)
        if keep:
            for offset, child in enumerate(list(element)):
                if not insertion:
                    _undelete(child)
                parent.insert(position + offset, deepcopy(child))
        parent.remove(element)

    package.save(str(path))


def _undelete(run: Any) -> None:
    from research_better.edit.word import XML_SPACE

    for node in list(run.iter(f"{W}delText")):
        replacement = run.makeelement(f"{W}t", {})
        replacement.text = node.text
        replacement.set(XML_SPACE, "preserve")
        node.getparent().replace(node, replacement)
